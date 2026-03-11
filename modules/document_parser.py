"""
Parser per diversi formati di documento
"""

import re
from pathlib import Path
from typing import Dict, List, Tuple
import docx
import docx2txt
import markdown


class DocumentParser:
    """Parser per documenti DOCX, LaTeX e Markdown"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.tex', '.md', '.markdown']
    
    def parse(self, file_path: Path) -> Dict[str, any]:
        """
        Parse il documento e restituisce contenuto strutturato
        
        Returns:
            Dict con 'content', 'title', 'sections', 'metadata'
        """
        suffix = file_path.suffix.lower()
        
        if suffix == '.docx':
            return self._parse_docx(file_path)
        elif suffix == '.tex':
            return self._parse_latex(file_path)
        elif suffix in ['.md', '.markdown']:
            return self._parse_markdown(file_path)
        else:
            raise ValueError(f"Formato non supportato: {suffix}")
    
    def _parse_docx(self, file_path: Path) -> Dict:
        """Parse documento DOCX"""
        doc = docx.Document(file_path)
        
        # Estrai testo
        text = docx2txt.process(str(file_path))
        
        # Estrai paragrafi strutturati
        sections = []
        current_section = {"title": "", "content": []}
        
        for para in doc.paragraphs:
            # Identifica titoli (basato su stile o font size)
            if para.style.name.startswith('Heading'):
                if current_section["content"]:
                    sections.append(current_section)
                current_section = {"title": para.text, "content": []}
            else:
                if para.text.strip():
                    current_section["content"].append(para.text)
        
        if current_section["content"]:
            sections.append(current_section)
        
        # Estrai titolo (primo paragrafo grande o primo heading)
        title = doc.paragraphs[0].text if doc.paragraphs else "Documento"
        
        return {
            "content": text,
            "title": title,
            "sections": sections,
            "metadata": {
                "format": "docx",
                "paragraphs": len(doc.paragraphs)
            }
        }
    
    def _parse_latex(self, file_path: Path) -> Dict:
        """Parse documento LaTeX"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Estrai titolo
        title_match = re.search(r'\\title\{([^}]+)\}', content)
        title = title_match.group(1) if title_match else "Documento LaTeX"
        
        # Estrai sezioni
        sections = []
        section_pattern = r'\\section\{([^}]+)\}(.*?)(?=\\section|\\end\{document\}|$)'
        
        for match in re.finditer(section_pattern, content, re.DOTALL):
            section_title = match.group(1)
            section_content = match.group(2)
            
            # Rimuovi comandi LaTeX comuni
            cleaned_content = re.sub(r'\\[a-zA-Z]+\{?', '', section_content)
            cleaned_content = re.sub(r'[{}]', '', cleaned_content)
            
            sections.append({
                "title": section_title,
                "content": [p.strip() for p in cleaned_content.split('\n\n') if p.strip()]
            })
        
        # Rimuovi comandi LaTeX dal contenuto completo
        cleaned_content = re.sub(r'\\[a-zA-Z]+(\[.*?\])?\{?', '', content)
        cleaned_content = re.sub(r'[{}%]', '', cleaned_content)
        
        return {
            "content": cleaned_content,
            "title": title,
            "sections": sections,
            "metadata": {
                "format": "latex",
                "sections_count": len(sections)
            }
        }
    
    def _parse_markdown(self, file_path: Path) -> Dict:
        """Parse documento Markdown"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Converti markdown in HTML per analisi
        md = markdown.Markdown(extensions=['meta', 'tables', 'fenced_code'])
        html = md.convert(content)
        
        # Estrai titolo (primo H1)
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        title = title_match.group(1) if title_match else "Documento Markdown"
        
        # Estrai sezioni (H2)
        sections = []
        section_pattern = r'^##\s+(.+)$\n(.*?)(?=^##\s+|$)'
        
        for match in re.finditer(section_pattern, content, re.MULTILINE | re.DOTALL):
            section_title = match.group(1)
            section_content = match.group(2)
            
            # Pulisci il contenuto
            paragraphs = [p.strip() for p in section_content.split('\n\n') if p.strip()]
            
            sections.append({
                "title": section_title,
                "content": paragraphs
            })
        
        return {
            "content": content,
            "title": title,
            "sections": sections,
            "metadata": {
                "format": "markdown",
                "sections_count": len(sections)
            }
        }
    
    def count_words(self, text: str) -> int:
        """Conta le parole in un testo"""
        return len(re.findall(r'\b\w+\b', text))
